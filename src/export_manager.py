"""
Export Manager

Export query results to PDF, Word, and Markdown formats
with full Unicode/Turkish character support
"""

import os
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
from loguru import logger
import re

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    REPORTLAB_AVAILABLE = True
    
    # Register Unicode fonts for Turkish support
    UNICODE_FONT_REGISTERED = False
    UNICODE_FONT_NAME = 'DejaVuSans'
    UNICODE_FONT_BOLD = 'DejaVuSans-Bold'
    
    def register_unicode_fonts():
        """Register Unicode-compatible fonts for Turkish character support"""
        global UNICODE_FONT_REGISTERED, UNICODE_FONT_NAME, UNICODE_FONT_BOLD
        
        if UNICODE_FONT_REGISTERED:
            return True
        
        # Possible font locations
        font_paths = [
            # Windows system fonts
            'C:/Windows/Fonts/DejaVuSans.ttf',
            'C:/Windows/Fonts/DejaVuSans-Bold.ttf',
            'C:/Windows/Fonts/arial.ttf',
            'C:/Windows/Fonts/arialbd.ttf',
            'C:/Windows/Fonts/seguiemj.ttf',  # Segoe UI Emoji for emojis
            # Linux
            '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
            '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
            # Project local fonts
            os.path.join(os.path.dirname(__file__), '..', 'fonts', 'DejaVuSans.ttf'),
            os.path.join(os.path.dirname(__file__), '..', 'fonts', 'DejaVuSans-Bold.ttf'),
        ]
        
        # Try to find and register DejaVu Sans (best Unicode support)
        dejavu_regular = None
        dejavu_bold = None
        
        for path in font_paths:
            if os.path.exists(path):
                if 'DejaVuSans-Bold' in path or 'arialbd' in path:
                    dejavu_bold = path
                elif 'DejaVuSans' in path or 'arial' in path.lower():
                    dejavu_regular = path
        
        # Register fonts
        try:
            if dejavu_regular:
                pdfmetrics.registerFont(TTFont('Turkish', dejavu_regular))
                UNICODE_FONT_NAME = 'Turkish'
                logger.info(f"✅ Registered Turkish font: {dejavu_regular}")
            else:
                # Fallback: Use Arial if available on Windows
                arial_path = 'C:/Windows/Fonts/arial.ttf'
                if os.path.exists(arial_path):
                    pdfmetrics.registerFont(TTFont('Turkish', arial_path))
                    UNICODE_FONT_NAME = 'Turkish'
                    logger.info(f"✅ Registered Turkish font (Arial): {arial_path}")
            
            if dejavu_bold:
                pdfmetrics.registerFont(TTFont('Turkish-Bold', dejavu_bold))
                UNICODE_FONT_BOLD = 'Turkish-Bold'
                logger.info(f"✅ Registered Turkish bold font: {dejavu_bold}")
            else:
                # Fallback: Use Arial Bold if available
                arial_bold_path = 'C:/Windows/Fonts/arialbd.ttf'
                if os.path.exists(arial_bold_path):
                    pdfmetrics.registerFont(TTFont('Turkish-Bold', arial_bold_path))
                    UNICODE_FONT_BOLD = 'Turkish-Bold'
                    logger.info(f"✅ Registered Turkish bold font (Arial): {arial_bold_path}")
            
            UNICODE_FONT_REGISTERED = True
            return True
            
        except Exception as e:
            logger.warning(f"⚠️ Could not register Unicode fonts: {e}")
            UNICODE_FONT_NAME = 'Helvetica'
            UNICODE_FONT_BOLD = 'Helvetica-Bold'
            return False
    
    # Try to register fonts at module load
    register_unicode_fonts()
    
except ImportError:
    REPORTLAB_AVAILABLE = False
    UNICODE_FONT_NAME = 'Helvetica'
    UNICODE_FONT_BOLD = 'Helvetica-Bold'
    logger.warning("⚠️ reportlab not installed - PDF export disabled")

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("⚠️ python-docx not installed - Word export disabled")


class ExportManager:
    """
    Export query results to various formats
    
    Features:
    - PDF export with Turkish support
    - Word document export
    - Markdown export
    """
    
    def __init__(self, export_dir: str = "exports"):
        """Initialize export manager"""
        self.export_dir = Path(export_dir)
        self.export_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"✅ Export Manager initialized at {export_dir}")
    
    def export_to_markdown(self, query: str, response: str, sources: List[Dict], 
                          metadata: Dict = None) -> str:
        """
        Export to Markdown
        
        Args:
            query: User query
            response: System response
            sources: Source documents
            metadata: Additional metadata
            
        Returns:
            Path to created file
        """
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"query_result_{timestamp}.md"
        filepath = self.export_dir / filename
        
        with open(filepath, 'w', encoding='utf-8') as f:
            # Header
            f.write("# PyRAG Query Result\n\n")
            f.write(f"**Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write("---\n\n")
            
            # Query
            f.write("## 📝 Query\n\n")
            f.write(f"{query}\n\n")
            f.write("---\n\n")
            
            # Response
            f.write("## 💡 Response\n\n")
            f.write(f"{response}\n\n")
            f.write("---\n\n")
            
            # Sources
            if sources:
                f.write("## 📚 Sources\n\n")
                for i, source in enumerate(sources, 1):
                    f.write(f"### Source {i}\n\n")
                    f.write(f"**Content:** {source.get('content', 'N/A')}\n\n")
                    
                    if 'metadata' in source:
                        meta = source['metadata']
                        f.write(f"**Document:** {meta.get('document_name', 'N/A')}\n\n")
                        f.write(f"**Section:** {meta.get('section_title', 'N/A')}\n\n")
                        f.write(f"**Page:** {meta.get('page_number', 'N/A')}\n\n")
                    
                    f.write("---\n\n")
            
            # Metadata
            if metadata:
                f.write("## ⚙️ Metadata\n\n")
                f.write(f"```json\n{metadata}\n```\n\n")
        
        logger.info(f"📄 Markdown exported: {filepath}")
        return str(filepath)
    
    def export_to_pdf(self, query: str, response: str, sources: List[Dict], 
                     metadata: Dict = None) -> Optional[str]:
        """
        Export to PDF with professional Markdown rendering
        
        Args:
            query: User query
            response: System response (can contain Markdown)
            sources: Source documents
            metadata: Additional metadata
            
        Returns:
            Path to created file or None if failed
        """
        if not REPORTLAB_AVAILABLE:
            logger.error("❌ PDF export not available - install reportlab")
            return None
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"query_result_{timestamp}.pdf"
        filepath = self.export_dir / filename
        
        try:
            # Create PDF
            doc = SimpleDocTemplate(str(filepath), pagesize=A4,
                                   rightMargin=2*cm, leftMargin=2*cm,
                                   topMargin=2*cm, bottomMargin=2*cm)
            
            # Container for elements
            story = []
            
            # Create enhanced styles
            styles = self._create_pdf_styles()
            
            # Title
            story.append(Paragraph("PyRAG Query Result", styles['title']))
            story.append(Paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['normal']))
            story.append(Spacer(1, 0.5*cm))
            
            # Query section
            story.append(Paragraph("❓ Query", styles['h1']))
            story.append(Paragraph(query, styles['normal']))
            story.append(Spacer(1, 0.5*cm))
            
            # Response section with Markdown parsing
            story.append(Paragraph("💡 Response", styles['h1']))
            story.append(Spacer(1, 0.3*cm))
            
            # Parse and render Markdown response
            self._render_markdown_to_pdf(response, story, styles)
            story.append(Spacer(1, 0.5*cm))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"📕 PDF exported: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"❌ PDF export failed: {e}")
            import traceback
            logger.error(traceback.format_exc())
            return None
    
    def _create_pdf_styles(self) -> Dict:
        """Create enhanced PDF styles with Turkish character support"""
        styles = getSampleStyleSheet()
        
        # Use Turkish font if available
        font_name = UNICODE_FONT_NAME if REPORTLAB_AVAILABLE else 'Helvetica'
        font_bold = UNICODE_FONT_BOLD if REPORTLAB_AVAILABLE else 'Helvetica-Bold'
        
        custom_styles = {
            'title': ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontName=font_bold,
                fontSize=22,
                textColor=colors.HexColor('#1E3A8A'),
                spaceAfter=20,
                alignment=1,  # Center
                borderWidth=2,
                borderColor=colors.HexColor('#3B82F6'),
                borderPadding=10,
                backColor=colors.HexColor('#EFF6FF')
            ),
            'h1': ParagraphStyle(
                'CustomH1',
                parent=styles['Heading1'],
                fontName=font_bold,
                fontSize=16,
                textColor=colors.HexColor('#1E40AF'),
                spaceAfter=12,
                spaceBefore=18,
                borderWidth=1,
                borderColor=colors.HexColor('#BFDBFE'),
                borderPadding=8,
                backColor=colors.HexColor('#EFF6FF')
            ),
            'h2': ParagraphStyle(
                'CustomH2',
                parent=styles['Heading2'],
                fontName=font_bold,
                fontSize=14,
                textColor=colors.HexColor('#2563EB'),
                spaceAfter=10,
                spaceBefore=14,
                leftIndent=5
            ),
            'h3': ParagraphStyle(
                'CustomH3',
                parent=styles['Heading3'],
                fontName=font_bold,
                fontSize=12,
                textColor=colors.HexColor('#059669'),
                spaceAfter=8,
                spaceBefore=12,
                leftIndent=10
            ),
            'normal': ParagraphStyle(
                'CustomNormal',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=10,
                spaceAfter=8,
                leading=16,
                textColor=colors.HexColor('#1F2937')
            ),
            'bullet': ParagraphStyle(
                'CustomBullet',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=10,
                leftIndent=25,
                bulletIndent=15,
                spaceAfter=5,
                leading=15,
                textColor=colors.HexColor('#374151')
            ),
            'warning': ParagraphStyle(
                'CustomWarning',
                parent=styles['Normal'],
                fontName=font_bold,
                fontSize=10,
                textColor=colors.HexColor('#B91C1C'),
                backColor=colors.HexColor('#FEF2F2'),
                borderWidth=2,
                borderColor=colors.HexColor('#F87171'),
                borderPadding=10,
                spaceAfter=10,
                spaceBefore=10
            ),
            'code': ParagraphStyle(
                'CustomCode',
                parent=styles['Normal'],
                fontName='Courier',
                fontSize=9,
                textColor=colors.HexColor('#BE185D'),
                backColor=colors.HexColor('#FDF2F8'),
                borderWidth=1,
                borderColor=colors.HexColor('#FBCFE8'),
                borderPadding=8,
                spaceAfter=10,
                leftIndent=15
            ),
            'source_ref': ParagraphStyle(
                'SourceRef',
                parent=styles['Normal'],
                fontName=font_name,
                fontSize=9,
                textColor=colors.HexColor('#6B7280'),
                leftIndent=15,
                spaceAfter=4,
                backColor=colors.HexColor('#F9FAFB'),
                borderPadding=5
            ),
            'highlight': ParagraphStyle(
                'Highlight',
                parent=styles['Normal'],
                fontName=font_bold,
                fontSize=11,
                textColor=colors.HexColor('#1E40AF'),
                backColor=colors.HexColor('#DBEAFE'),
                borderPadding=8,
                spaceAfter=10
            )
        }
        
        return custom_styles
    
    def _render_markdown_to_pdf(self, text: str, story: list, styles: Dict):
        """
        Parse Markdown text and render to PDF elements
        
        Args:
            text: Markdown formatted text
            story: ReportLab story list to append to
            styles: Style dictionary
        """
        import re
        
        lines = text.split('\n')
        in_table = False
        table_data = []
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Table detection
            if '|' in line and line.strip().startswith('|'):
                cells = [c.strip() for c in line.strip().strip('|').split('|')]
                
                # Skip separator lines
                if all(set(c.strip()).issubset({'-', ':', ' '}) for c in cells if c.strip()):
                    i += 1
                    continue
                
                if not in_table:
                    in_table = True
                    table_data = []
                
                table_data.append(cells)
                
                # Check if next line is end of table
                if i + 1 >= len(lines) or '|' not in lines[i + 1]:
                    # Render table
                    self._render_table_to_pdf(table_data, story, styles)
                    in_table = False
                    table_data = []
                
                i += 1
                continue
            
            # Headers
            if line.startswith('### '):
                story.append(Paragraph(line[4:], styles['h3']))
            elif line.startswith('## '):
                story.append(Paragraph(line[3:], styles['h2']))
            elif line.startswith('# '):
                story.append(Paragraph(line[2:], styles['h1']))
            
            # Bullet points
            elif re.match(r'^\s*[-*•]\s+', line):
                content = re.sub(r'^\s*[-*•]\s+', '', line)
                # Remove ** markers
                content = content.replace('**', '')
                # Check for "Label:" pattern and make label bold
                colon_match = re.match(r'^([^:]+):\s*(.*)$', content)
                if colon_match:
                    label = colon_match.group(1)
                    rest = colon_match.group(2)
                    content = f"<b>{label}:</b> {self._process_inline_formatting(rest)}"
                else:
                    content = self._process_inline_formatting(content)
                story.append(Paragraph(f"• {content}", styles['bullet']))
            
            # Numbered lists
            elif re.match(r'^\s*\d+\.\s+', line):
                content = re.sub(r'^\s*\d+\.\s+', '', line)
                num_match = re.match(r'^\s*(\d+)\.', line)
                num = num_match.group(1) if num_match else '•'
                content = self._process_inline_formatting(content)
                story.append(Paragraph(f"{num}. {content}", styles['bullet']))
            
            # Warning/Important
            elif any(marker in line.lower() for marker in ['warning:', 'important:', '⚠', '❗', 'note:']):
                story.append(Paragraph(self._process_inline_formatting(line), styles['warning']))
            
            # Source references
            elif any(marker in line for marker in ['📚', '📄', 'Source:', '[Source', '(Page', 'Reference:']):
                story.append(Paragraph(line, styles['source_ref']))
            
            # Regular paragraph
            elif line.strip():
                content = self._process_inline_formatting(line)
                story.append(Paragraph(content, styles['normal']))
            
            # Empty line = spacer
            else:
                story.append(Spacer(1, 0.2*cm))
            
            i += 1
    
    def _process_inline_formatting(self, text: str) -> str:
        """Convert Markdown inline formatting to ReportLab XML tags"""
        import re
        
        # Note: ** are already cleaned in bullet processing
        # Only process * for italic
        
        # Italic: *text* -> <i>text</i> (but not **)
        text = re.sub(r'(?<!\*)\*([^*]+?)\*(?!\*)', r'<i>\1</i>', text)
        
        # Code: `text` -> <font face="Courier" color="#DC2626">text</font>
        text = re.sub(r'`(.+?)`', r'<font face="Courier" color="#DC2626">\1</font>', text)
        
        return text
    
    def _render_table_to_pdf(self, table_data: List[List[str]], story: list, styles: Dict):
        """Render a Markdown table to PDF with professional styling"""
        if not table_data:
            return
        
        # Use Turkish font
        font_name = UNICODE_FONT_NAME if REPORTLAB_AVAILABLE else 'Helvetica'
        font_bold = UNICODE_FONT_BOLD if REPORTLAB_AVAILABLE else 'Helvetica-Bold'
        
        # Calculate column widths - use more page width
        num_cols = max(len(row) for row in table_data)
        available_width = 17 * cm  # Use more of the page width
        
        # Smart column width distribution based on typical content
        if num_cols == 5:  # Common case: Dedektör, Alan, Mesafe, Mesafe, Referans
            col_widths = [4*cm, 3*cm, 3*cm, 3*cm, 4*cm]
        else:
            col_width = available_width / num_cols
            col_widths = [col_width] * num_cols
        
        # Clean and pad data with better wrapping
        clean_data = []
        for row_idx, row in enumerate(table_data):
            clean_row = []
            for col_idx, cell in enumerate(row):
                cell_text = self._process_inline_formatting(cell)
                
                # For header row, wrap at 20 chars; for data, wrap at 25
                if row_idx == 0 and len(cell_text) > 20:
                    cell_text = self._wrap_header_text(cell_text, max_chars=20)
                elif row_idx > 0 and len(cell_text) > 25:
                    cell_text = self._wrap_header_text(cell_text, max_chars=25)
                
                clean_row.append(cell_text)
            
            # Pad row if needed
            while len(clean_row) < num_cols:
                clean_row.append('')
            clean_data.append(clean_row[:num_cols])
        
        # Create table
        table = Table(clean_data, colWidths=col_widths)
        
        # Style table with professional look
        style_commands = [
            # Header row - Professional blue
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E40AF')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), font_bold),
            ('FONTSIZE', (0, 0), (-1, 0), 8),  # Smaller font for headers
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            
            # Data rows - Turkish font
            ('FONTNAME', (0, 1), (-1, -1), font_name),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.HexColor('#F8FAFC'), colors.white]),
            ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#1F2937')),
            
            # Grid - Subtle borders
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#94A3B8')),
            ('BOX', (0, 0), (-1, -1), 1.5, colors.HexColor('#1E40AF')),
            ('TOPPADDING', (0, 0), (-1, -1), 6),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
            ('LEFTPADDING', (0, 0), (-1, -1), 6),
            ('RIGHTPADDING', (0, 0), (-1, -1), 6),
        ]
        
        table.setStyle(TableStyle(style_commands))
        story.append(Spacer(1, 0.3*cm))
        story.append(table)
        story.append(Spacer(1, 0.4*cm))
    
    def _wrap_header_text(self, text: str, max_chars: int = 12) -> str:
        """
        Wrap long header text to multiple lines
        
        Args:
            text: Header text to wrap
            max_chars: Maximum characters per line
            
        Returns:
            Text with line breaks for wrapping
        """
        if len(text) <= max_chars:
            return text
        
        words = text.split()
        lines = []
        current_line = ""
        
        for word in words:
            if not current_line:
                current_line = word
            elif len(current_line) + 1 + len(word) <= max_chars:
                current_line += " " + word
            else:
                lines.append(current_line)
                current_line = word
        
        if current_line:
            lines.append(current_line)
        
        return "\n".join(lines)
    
    def export_to_word(self, query: str, response: str, sources: List[Dict], 
                      metadata: Dict = None) -> Optional[str]:
        """
        Export to Word
        
        Args:
            query: User query
            response: System response
            sources: Source documents
            metadata: Additional metadata
            
        Returns:
            Path to created file or None if failed
        """
        if not DOCX_AVAILABLE:
            logger.error("❌ Word export not available - install python-docx")
            return None
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"query_result_{timestamp}.docx"
        filepath = self.export_dir / filename
        
        try:
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading('PyRAG Query Result', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()  # Space
            
            # Query
            doc.add_heading('📝 Query', level=1)
            doc.add_paragraph(query)
            
            # Response
            doc.add_heading('💡 Response', level=1)
            doc.add_paragraph(response)
            
            # Sources
            if sources:
                doc.add_heading('📚 Sources', level=1)
                for i, source in enumerate(sources, 1):
                    doc.add_heading(f'Source {i}', level=2)
                    
                    # Content
                    content = source.get('content', 'N/A')[:500]  # Limit length
                    doc.add_paragraph(content)
                    
                    # Metadata
                    if 'metadata' in source:
                        meta = source['metadata']
                        meta_para = doc.add_paragraph()
                        meta_para.add_run('Document: ').bold = True
                        meta_para.add_run(f"{meta.get('document_name', 'N/A')}\n")
                        meta_para.add_run('Section: ').bold = True
                        meta_para.add_run(f"{meta.get('section_title', 'N/A')}\n")
                        meta_para.add_run('Page: ').bold = True
                        meta_para.add_run(f"{meta.get('page_number', 'N/A')}\n")
            
            # Save
            doc.save(str(filepath))
            
            logger.info(f"📘 Word document exported: {filepath}")
            return str(filepath)
            
        except Exception as e:
            logger.error(f"❌ Word export failed: {e}")
            return None
    
    def export_all_formats(self, query: str, response: str, sources: List[Dict], 
                          metadata: Dict = None) -> Dict[str, Optional[str]]:
        """
        Export to all available formats
        
        Returns:
            Dict with format -> filepath mapping
        """
        results = {}
        
        results['markdown'] = self.export_to_markdown(query, response, sources, metadata)
        results['pdf'] = self.export_to_pdf(query, response, sources, metadata)
        results['word'] = self.export_to_word(query, response, sources, metadata)
        
        return results


# Singleton instance
_export_manager = None

def get_export_manager() -> ExportManager:
    """Get or create export manager singleton"""
    global _export_manager
    
    if _export_manager is None:
        _export_manager = ExportManager()
    
    return _export_manager
